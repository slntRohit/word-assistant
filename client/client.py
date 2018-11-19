from docx import Document
import json
import websocket
import logging
from uuid import getnode


WS_URL = "wss://word-assistant.herokuapp.com/connect"

# TODO: Add more functionality
class DocumentWriter:

    def __init__(self, docname="generated.docx"):
        self._current_para = None
        self.error_code = 0
        self._document = Document()
        self._docname = docname

    def add_text(self, text):
        if self._current_para is None:
            self.add_paragraph(text)
        else:
            self._current_para.add_run(text=text)

    def add_heading(self, text):
        self._current_para = self._document.add_heading(text=text)

    def add_paragraph(self, text):
        self._current_para = self._document.add_paragraph(text=text)

    def save_document(self):
        self._document.save(self._docname)


class WebSocketHandler:

    def __init__(self):
        self.doc = DocumentWriter()

    def parse_message(self, command):
        json_data = json.loads(command)
        query = str(json_data['queryText'])
        parameters = json_data['parameters']['text']
        command = query.split()[0]
        return command, parameters


main = WebSocketHandler()


def on_message(ws, message):
    print("command", message)
    command, parameters = main.parse_message(message)
    if command == 'type':
        try:
            print("adding", parameters)
            main.doc.add_text(parameters)
            main.doc.save_document()
            ws.send("Message processed")
        except Exception as e:
            logging.error('Failed to add to document: ' + str(e))
            ws.send("Message not processed")

    if command == 'create':
        try:
            print("creating", parameters)
            # doc.add_text(parameters)
            main.doc = DocumentWriter(docname=parameters)
            main.doc.save_document()
            ws.send("Message processed")
        except Exception as e:
            logging.error('Failed to add to document: ' + str(e))
            ws.send("Message not processed")


def on_error(ws, error):
    print(error)


def on_close(ws):
    print("### closed ###")
    connect()


def on_open(ws):
    ws.send(str(getnode()))
    print("connected with id: " + str(getnode()))


def connect():
    websocket.enableTrace(True)
    ws = websocket.WebSocketApp(WS_URL,
                                on_message=on_message,
                                on_error=on_error,
                                on_close=on_close)
    ws.on_open = on_open
    ws.run_forever()


if __name__ == "__main__":
    connect()


#
# iter = 0
#
# client_busy = False
# client_online = True
# message_from_cloud = []
#
# while True:
#     if not client_online:
#         ws.send("client offline")
#         print("client offline")
#
#     if client_busy:
#         ws.send("client busy")
#         message_clientID, message_filename, message_command, message_value, message_count = parse_message(message_from_cloud)
#         document_name = 'outputs/' + message_filename + '.docx'
#         doc = DocumentWriter(docname=document_name)
#
#         #Create a file
#         if message_command == 'create':
#             try:
#                 doc.add_paragraph(message_value)
#                 doc.save_document()
#                 ws.send("Message %s processed, 0" % message_count)
#                 client_busy = False
#             except:
#                 ws.send("Message %s processed, 1" % message_count)
#                 client_busy = False
#
#         # Edit an existing File
#         if message_command == 'edit':
#             try:
#                 doc.add_paragraph(message_value)
#                 doc.save_document()
#                 ws.send("Message %s processed, 0" % message_count)
#                 client_busy = False
#             except:
#                 ws.send("Message %s processed, 1" % message_count)
#                 client_busy = False
#
#         # Delete a file
#         if message_command == 'delete':
#             try:
#                 os.remove(document_name)
#                 ws.send("Message %s processed, 0" % message_count)
#                 client_busy = False
#             except OSError:
#                 ws.send("Message %s processed, 1" % message_count)
#                 client_busy = False
#
#     if not client_busy and client_online:
#         print("Sending 'ready to receive from client'")
#         ws.send("client_online and not busy")
#         print("Sent to cloud")
#         print("Waiting for data from cloud")
#         message_from_cloud = ws.recv()
#         print("Message Received from cloud")
#         print(message_from_cloud)
#         client_busy = True
